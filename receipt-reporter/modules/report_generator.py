"""
Report Generator Module — Excel, Word, PDF 출장 보고서 생성
"""
import os
from datetime import datetime
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from .categorizer import get_category_group, check_qualified_receipt, calculate_summary, is_vat_deductible


# 공통 설정
CATEGORY_LABELS = {
    "T01": "항공료", "T02": "철도/버스", "T03": "택시", "T04": "렌터카",
    "T05": "자가용/주유", "T06": "대중교통",
    "A01": "호텔", "A02": "기타 숙박",
    "M01": "식비", "M02": "음료/간식",
    "E01": "접대비", "E02": "선물/기념품",
    "C01": "통신비",
    "R01": "등록/참가비",
    "S01": "업무용품",
    "O01": "비자/여권/보험", "O99": "기타",
}


def _format_amount(amount, currency="KRW"):
    """금액 포맷"""
    if amount is None:
        return "-"
    if currency == "KRW":
        return f"₩{amount:,.0f}"
    return f"{amount:,.2f}"


class ExcelReportGenerator:
    """Excel(.xlsx) 보고서 생성기"""

    def generate(self, trip_info: dict, receipts: list, output_path: str):
        wb = Workbook()

        self._create_overview_sheet(wb, trip_info, receipts)
        self._create_detail_sheet(wb, trip_info, receipts)
        self._create_summary_sheet(wb, trip_info, receipts)
        self._create_payment_sheet(wb, trip_info, receipts)
        self._create_tax_sheet(wb, trip_info, receipts)
        self._create_approval_sheet(wb, trip_info)

        wb.save(output_path)
        return output_path

    def _create_overview_sheet(self, wb, trip_info, receipts):
        ws = wb.active
        ws.title = "출장 개요"

        header_font = Font(name="맑은 고딕", size=16, bold=True)
        label_font = Font(name="맑은 고딕", size=11, bold=True)
        value_font = Font(name="맑은 고딕", size=11)
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )
        header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")

        ws.merge_cells("A1:F1")
        ws["A1"] = "출장 보고서"
        ws["A1"].font = Font(name="맑은 고딕", size=18, bold=True, color="FFFFFF")
        ws["A1"].fill = header_fill
        ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 40

        summary = calculate_summary(
            receipts,
            trip_info.get("trip_type", "domestic"),
            trip_info.get("exchange_rate", 1.0)
        )

        fields = [
            ("보고서 번호", trip_info.get("report_id", "")),
            ("작성일", datetime.now().strftime("%Y-%m-%d")),
            ("출장자", trip_info.get("employee_name", "")),
            ("소속/직급", f"{trip_info.get('department', '')} / {trip_info.get('position', '')}"),
            ("출장 유형", "해외 출장" if trip_info.get("trip_type") == "overseas" else "국내 출장"),
            ("출장 기간", f"{trip_info.get('start_date', '')} ~ {trip_info.get('end_date', '')}"),
            ("출장지", trip_info.get("destination", "")),
            ("방문처", trip_info.get("visit_company", "")),
            ("출장 목적", trip_info.get("purpose", "")),
            ("동행자", trip_info.get("attendees", "")),
            ("출장 결과", trip_info.get("result", "")),
            ("후속 조치", trip_info.get("follow_up", "")),
            ("총 비용", _format_amount(summary["total_amount"])),
            ("영수증 건수", f"{summary['total_count']}건"),
        ]

        for i, (label, value) in enumerate(fields, start=3):
            ws.cell(row=i, column=1, value=label).font = label_font
            ws.cell(row=i, column=1).border = thin_border
            ws.cell(row=i, column=1).fill = PatternFill(start_color="D6E4F0", fill_type="solid")
            ws.merge_cells(start_row=i, start_column=2, end_row=i, end_column=6)
            ws.cell(row=i, column=2, value=str(value)).font = value_font
            ws.cell(row=i, column=2).border = thin_border

        ws.column_dimensions["A"].width = 18
        for col in range(2, 7):
            ws.column_dimensions[get_column_letter(col)].width = 15

    def _create_detail_sheet(self, wb, trip_info, receipts):
        ws = wb.create_sheet("비용 상세")

        header_font = Font(name="맑은 고딕", size=10, bold=True, color="FFFFFF")
        data_font = Font(name="맑은 고딕", size=10)
        header_fill = PatternFill(start_color="1F4E79", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        is_overseas = trip_info.get("trip_type") == "overseas"
        exchange_rate = trip_info.get("exchange_rate", 1.0)

        if is_overseas:
            headers = ["No", "날짜", "카테고리", "사용처", "내역", "현지 금액", "통화", "적용 환율", "원화 금액", "결제수단", "비고"]
        else:
            headers = ["No", "날짜", "카테고리", "사용처", "내역", "공급가액", "부가세", "합계", "결제수단", "증빙유형", "비고"]

        for j, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=j, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border

        total = 0
        for i, r in enumerate(receipts, start=2):
            amount = r.get("total_amount", 0) or 0
            krw = amount if r.get("currency", "KRW") == "KRW" else round(amount * exchange_rate)
            total += krw

            cat_code = r.get("category", "O99")
            cat_label = CATEGORY_LABELS.get(cat_code, cat_code)

            if is_overseas:
                row_data = [
                    i - 1, r.get("date", ""), cat_label,
                    r.get("vendor_name", ""), r.get("description", ""),
                    amount, r.get("currency", ""), exchange_rate, krw,
                    r.get("payment_method", ""), r.get("remarks", ""),
                ]
            else:
                row_data = [
                    i - 1, r.get("date", ""), cat_label,
                    r.get("vendor_name", ""), r.get("description", ""),
                    r.get("supply_amount", ""), r.get("vat_amount", ""), amount,
                    r.get("payment_method", ""), r.get("receipt_type", ""), r.get("remarks", ""),
                ]

            for j, val in enumerate(row_data, 1):
                cell = ws.cell(row=i, column=j, value=val)
                cell.font = data_font
                cell.border = thin_border
                if isinstance(val, (int, float)) and j >= 6:
                    cell.number_format = '#,##0'

        # 합계 행
        total_row = len(receipts) + 2
        ws.cell(row=total_row, column=1, value="합계").font = Font(name="맑은 고딕", bold=True)
        total_col = 9 if is_overseas else 8
        ws.cell(row=total_row, column=total_col, value=total).font = Font(name="맑은 고딕", bold=True)
        ws.cell(row=total_row, column=total_col).number_format = '#,##0'

        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 14

    def _create_summary_sheet(self, wb, trip_info, receipts):
        ws = wb.create_sheet("카테고리 요약")

        summary = calculate_summary(
            receipts,
            trip_info.get("trip_type", "domestic"),
            trip_info.get("exchange_rate", 1.0)
        )

        header_font = Font(name="맑은 고딕", size=11, bold=True)

        ws.cell(row=1, column=1, value="카테고리별 지출 요약").font = Font(name="맑은 고딕", size=14, bold=True)

        ws.cell(row=3, column=1, value="카테고리").font = header_font
        ws.cell(row=3, column=2, value="금액 (원)").font = header_font
        ws.cell(row=3, column=3, value="비율").font = header_font

        row = 4
        for cat, amount in sorted(summary["category_totals"].items()):
            pct = round(amount / max(summary["total_amount"], 1) * 100, 1)
            ws.cell(row=row, column=1, value=cat)
            ws.cell(row=row, column=2, value=amount).number_format = '#,##0'
            ws.cell(row=row, column=3, value=f"{pct}%")
            row += 1

        ws.cell(row=row, column=1, value="합계").font = header_font
        ws.cell(row=row, column=2, value=summary["total_amount"]).number_format = '#,##0'

        row += 2
        ws.cell(row=row, column=1, value="세무 검증 요약").font = Font(name="맑은 고딕", size=14, bold=True)
        row += 1
        trip_type_label = "해외 출장 (적격증빙 면제)" if summary["trip_type"] == "overseas" else "국내 출장"
        tax_items = [
            ("출장 유형", trip_type_label),
            ("적격증빙 수취율", f"{summary['qualified_count']}/{summary['total_count']}건 ({summary['qualified_rate']}%)"),
            ("부가세 공제 가능액", _format_amount(summary["vat_deductible"])),
            ("부가세 불공제액", _format_amount(summary["vat_non_deductible"])),
            ("가산세 위험 건", f"{summary['risk_items_count']}건"),
        ]
        for label, value in tax_items:
            ws.cell(row=row, column=1, value=label).font = header_font
            ws.cell(row=row, column=2, value=value)
            row += 1

        ws.column_dimensions["A"].width = 22
        ws.column_dimensions["B"].width = 18
        ws.column_dimensions["C"].width = 10

    def _create_payment_sheet(self, wb, trip_info, receipts):
        """결제수단별 분석 시트 (#13 보완)"""
        ws = wb.create_sheet("결제수단 분석")
        header_font = Font(name="맑은 고딕", size=11, bold=True)
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )
        exchange_rates = trip_info.get("exchange_rates", {})
        default_rate = trip_info.get("exchange_rate", 1.0)

        ws.cell(row=1, column=1, value="결제수단별 지출 분석").font = Font(name="맑은 고딕", size=14, bold=True)

        headers = ["결제수단", "건수", "금액 (원)", "비율"]
        for j, h in enumerate(headers, 1):
            ws.cell(row=3, column=j, value=h).font = header_font
            ws.cell(row=3, column=j).border = thin_border

        payment_data = {}
        for r in receipts:
            pm = r.get("payment_method", "기타")
            amount = r.get("total_amount", 0) or 0
            currency = r.get("currency", "KRW")
            if currency == "KRW":
                krw = amount
            else:
                rate = exchange_rates.get(currency, default_rate)
                krw = round(amount * rate)
            if pm not in payment_data:
                payment_data[pm] = {"count": 0, "amount": 0}
            payment_data[pm]["count"] += 1
            payment_data[pm]["amount"] += krw

        total = sum(d["amount"] for d in payment_data.values())
        row = 4
        for pm, data in sorted(payment_data.items(), key=lambda x: -x[1]["amount"]):
            pct = round(data["amount"] / max(total, 1) * 100, 1)
            ws.cell(row=row, column=1, value=pm).border = thin_border
            ws.cell(row=row, column=2, value=data["count"]).border = thin_border
            ws.cell(row=row, column=3, value=data["amount"]).border = thin_border
            ws.cell(row=row, column=3).number_format = '#,##0'
            ws.cell(row=row, column=4, value=f"{pct}%").border = thin_border
            row += 1

        ws.cell(row=row, column=1, value="합계").font = header_font
        ws.cell(row=row, column=2, value=len(receipts))
        ws.cell(row=row, column=3, value=total).number_format = '#,##0'

        for col in range(1, 5):
            ws.column_dimensions[get_column_letter(col)].width = 18

    def _create_tax_sheet(self, wb, trip_info, receipts):
        """세무 검증 상세 시트 (#13 보완)"""
        ws = wb.create_sheet("세무 검증")
        header_font = Font(name="맑은 고딕", size=11, bold=True)
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )
        trip_type = trip_info.get("trip_type", "domestic")

        ws.cell(row=1, column=1, value="세무 검증 상세").font = Font(name="맑은 고딕", size=14, bold=True)

        # 영수증별 세무 검증 테이블
        headers = ["No", "날짜", "사용처", "금액", "증빙유형", "적격증빙", "VAT 공제", "비고"]
        for j, h in enumerate(headers, 1):
            ws.cell(row=3, column=j, value=h).font = header_font
            ws.cell(row=3, column=j).border = thin_border

        row = 4
        for i, r in enumerate(receipts):
            qual = check_qualified_receipt(r)
            vat_ok = is_vat_deductible(r, trip_type)
            ws.cell(row=row, column=1, value=i + 1).border = thin_border
            ws.cell(row=row, column=2, value=r.get("date", "")).border = thin_border
            ws.cell(row=row, column=3, value=r.get("vendor_name", "")).border = thin_border
            ws.cell(row=row, column=4, value=r.get("total_amount", 0) or 0).border = thin_border
            ws.cell(row=row, column=4).number_format = '#,##0'
            ws.cell(row=row, column=5, value=r.get("receipt_type", "")).border = thin_border
            ws.cell(row=row, column=6, value="O" if qual["is_qualified"] else "X").border = thin_border
            ws.cell(row=row, column=7, value="O" if vat_ok else "X").border = thin_border
            ws.cell(row=row, column=8, value=qual["message"]).border = thin_border

            # 위험 항목 빨간 배경
            if qual["risk_level"] == "warning":
                for col in range(1, 9):
                    ws.cell(row=row, column=col).fill = PatternFill(start_color="FFE0E0", fill_type="solid")
            row += 1

        # 접대비 상세 섹션
        entertainment = [r for r in receipts if r.get("category", "").startswith("E")]
        if entertainment:
            row += 1
            ws.cell(row=row, column=1, value="접대비 상세").font = Font(name="맑은 고딕", size=14, bold=True)
            row += 1
            ent_headers = ["No", "날짜", "사용처", "금액", "참석자", "비고"]
            for j, h in enumerate(ent_headers, 1):
                ws.cell(row=row, column=j, value=h).font = header_font
                ws.cell(row=row, column=j).border = thin_border
            row += 1
            for i, r in enumerate(entertainment):
                ws.cell(row=row, column=1, value=i + 1).border = thin_border
                ws.cell(row=row, column=2, value=r.get("date", "")).border = thin_border
                ws.cell(row=row, column=3, value=r.get("vendor_name", "")).border = thin_border
                ws.cell(row=row, column=4, value=r.get("total_amount", 0) or 0).border = thin_border
                ws.cell(row=row, column=4).number_format = '#,##0'
                ws.cell(row=row, column=5, value=r.get("participants", "")).border = thin_border
                ws.cell(row=row, column=6, value=r.get("remarks", "")).border = thin_border
                row += 1

        for col in range(1, 9):
            ws.column_dimensions[get_column_letter(col)].width = 16

    def _create_approval_sheet(self, wb, trip_info):
        """결재란 시트"""
        ws = wb.create_sheet("결재")
        header_font = Font(name="맑은 고딕", size=11, bold=True)
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        ws.cell(row=1, column=1, value="결재").font = Font(name="맑은 고딕", size=14, bold=True)

        headers = ["구분", "작성자", "팀장", "재무팀", "최종 승인"]
        for j, h in enumerate(headers, 1):
            cell = ws.cell(row=3, column=j, value=h)
            cell.font = header_font
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center")

        ws.cell(row=4, column=1, value="성명").font = header_font
        ws.cell(row=4, column=1).border = thin_border
        ws.cell(row=4, column=2, value=trip_info.get("employee_name", "")).border = thin_border
        for j in range(3, 6):
            ws.cell(row=4, column=j).border = thin_border

        ws.cell(row=5, column=1, value="서명").font = header_font
        ws.cell(row=5, column=1).border = thin_border
        ws.row_dimensions[5].height = 50
        for j in range(2, 6):
            ws.cell(row=5, column=j).border = thin_border

        ws.cell(row=6, column=1, value="날짜").font = header_font
        ws.cell(row=6, column=1).border = thin_border
        ws.cell(row=6, column=2, value=datetime.now().strftime("%Y-%m-%d")).border = thin_border
        for j in range(3, 6):
            ws.cell(row=6, column=j).border = thin_border

        for col in range(1, 6):
            ws.column_dimensions[get_column_letter(col)].width = 16


class WordReportGenerator:
    """Word(.docx) 보고서 생성기"""

    def _set_korean_font(self, doc):
        """Word 문서 전체에 한글 폰트 적용"""
        from docx.oxml.ns import qn
        style = doc.styles["Normal"]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(10)
        # eastAsia 폰트 설정 (한글 표시에 필수)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    def generate(self, trip_info: dict, receipts: list, output_path: str):
        doc = Document()
        self._set_korean_font(doc)

        self._add_title(doc, trip_info)
        self._add_overview(doc, trip_info, receipts)
        self._add_expense_detail(doc, trip_info, receipts)
        self._add_summary(doc, trip_info, receipts)
        self._add_approval(doc)

        doc.save(output_path)
        return output_path

    def _add_title(self, doc, trip_info):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("출장 비용 지출 보고서")
        run.font.size = Pt(18)
        run.font.bold = True
        doc.add_paragraph()

    def _add_overview(self, doc, trip_info, receipts):
        doc.add_heading("1. 출장 개요", level=2)

        summary = calculate_summary(
            receipts,
            trip_info.get("trip_type", "domestic"),
            trip_info.get("exchange_rate", 1.0)
        )

        table = doc.add_table(rows=10, cols=4, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        fields = [
            ("출장자", trip_info.get("employee_name", ""), "소속/직급", f"{trip_info.get('department', '')} / {trip_info.get('position', '')}"),
            ("출장 기간", f"{trip_info.get('start_date', '')} ~ {trip_info.get('end_date', '')}", "출장 유형", "해외" if trip_info.get("trip_type") == "overseas" else "국내"),
            ("출장지", trip_info.get("destination", ""), "방문처", trip_info.get("visit_company", "")),
            ("출장 목적", trip_info.get("purpose", ""), "", ""),
            ("동행자", trip_info.get("attendees", ""), "", ""),
            ("출장 결과", trip_info.get("result", ""), "", ""),
            ("후속 조치", trip_info.get("follow_up", ""), "", ""),
            ("총 비용", _format_amount(summary["total_amount"]), "영수증", f"{summary['total_count']}건"),
            ("적격증빙", f"{summary['qualified_rate']}%", "가산세 위험", f"{summary['risk_items_count']}건"),
            ("작성일", datetime.now().strftime("%Y-%m-%d"), "보고서 번호", trip_info.get("report_id", "")),
        ]

        for i, (l1, v1, l2, v2) in enumerate(fields):
            row = table.rows[i]
            row.cells[0].text = l1
            row.cells[1].text = str(v1)
            row.cells[2].text = l2
            row.cells[3].text = str(v2)
            for cell in [row.cells[0], row.cells[2]]:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

        doc.add_paragraph()

    def _add_expense_detail(self, doc, trip_info, receipts):
        doc.add_heading("2. 비용 지출 상세", level=2)

        is_overseas = trip_info.get("trip_type") == "overseas"
        exchange_rate = trip_info.get("exchange_rate", 1.0)

        if is_overseas:
            headers = ["No", "날짜", "카테고리", "사용처", "현지금액", "환율", "원화금액"]
        else:
            headers = ["No", "날짜", "카테고리", "사용처", "공급가액", "부가세", "합계"]

        table = doc.add_table(rows=1 + len(receipts) + 1, cols=len(headers), style="Table Grid")

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True

        total = 0
        for i, r in enumerate(receipts):
            amount = r.get("total_amount", 0) or 0
            krw = amount if r.get("currency", "KRW") == "KRW" else round(amount * exchange_rate)
            total += krw

            cat_label = CATEGORY_LABELS.get(r.get("category", "O99"), "기타")
            row = table.rows[i + 1]

            if is_overseas:
                vals = [str(i + 1), r.get("date", ""), cat_label, r.get("vendor_name", ""),
                        f"{amount:,.0f}", f"{exchange_rate:,.2f}", f"{krw:,.0f}"]
            else:
                vals = [str(i + 1), r.get("date", ""), cat_label, r.get("vendor_name", ""),
                        f"{r.get('supply_amount', 0) or 0:,.0f}", f"{r.get('vat_amount', 0) or 0:,.0f}", f"{amount:,.0f}"]

            for j, v in enumerate(vals):
                row.cells[j].text = v

        # 합계 행
        total_row = table.rows[-1]
        total_row.cells[0].text = ""
        total_row.cells[1].text = ""
        total_row.cells[2].text = ""
        total_row.cells[3].text = "합계"
        total_row.cells[-1].text = f"{total:,.0f}"
        for p in total_row.cells[3].paragraphs:
            for r in p.runs:
                r.font.bold = True
        for p in total_row.cells[-1].paragraphs:
            for r in p.runs:
                r.font.bold = True

        doc.add_paragraph()

    def _add_summary(self, doc, trip_info, receipts):
        doc.add_heading("3. 카테고리별 요약", level=2)

        summary = calculate_summary(
            receipts,
            trip_info.get("trip_type", "domestic"),
            trip_info.get("exchange_rate", 1.0)
        )

        table = doc.add_table(rows=1 + len(summary["category_totals"]) + 1, cols=3, style="Table Grid")
        table.rows[0].cells[0].text = "카테고리"
        table.rows[0].cells[1].text = "금액"
        table.rows[0].cells[2].text = "비율"

        row_idx = 1
        for cat, amount in sorted(summary["category_totals"].items()):
            pct = round(amount / max(summary["total_amount"], 1) * 100, 1)
            table.rows[row_idx].cells[0].text = cat
            table.rows[row_idx].cells[1].text = f"₩{amount:,.0f}"
            table.rows[row_idx].cells[2].text = f"{pct}%"
            row_idx += 1

        table.rows[row_idx].cells[0].text = "합계"
        table.rows[row_idx].cells[1].text = f"₩{summary['total_amount']:,.0f}"

        doc.add_paragraph()

    def _add_approval(self, doc):
        doc.add_heading("4. 결재", level=2)

        table = doc.add_table(rows=3, cols=4, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["구분", "작성자", "팀장", "최종 승인"]
        for j, h in enumerate(headers):
            table.rows[0].cells[j].text = h

        table.rows[1].cells[0].text = "성명"
        table.rows[2].cells[0].text = "서명/날짜"

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


class PDFReportGenerator:
    """PDF 보고서 생성기"""

    def __init__(self):
        self._register_font()

    def _register_font(self):
        """한글 폰트 등록 (맑은 고딕 우선)"""
        self.font_registered = False
        # TTF 파일만 사용 (TTC는 ReportLab에서 인덱스 지정 필요)
        font_candidates = [
            ("C:/Windows/Fonts/malgun.ttf", "Korean"),
            ("C:/Windows/Fonts/malgunbd.ttf", "KoreanBold"),
        ]
        for fp, name in font_candidates:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont(name, fp))
                    if name == "Korean":
                        self.font_registered = True
                except Exception:
                    continue

        if not self.font_registered:
            # 폴백: 시스템의 다른 TTF 한글 폰트 탐색
            import glob
            for fp in glob.glob("C:/Windows/Fonts/*.ttf"):
                try:
                    pdfmetrics.registerFont(TTFont("Korean", fp))
                    # 한글 테스트
                    from reportlab.pdfbase.pdfmetrics import getFont
                    getFont("Korean")
                    self.font_registered = True
                    break
                except Exception:
                    continue

    def generate(self, trip_info: dict, receipts: list, output_path: str):
        doc = SimpleDocTemplate(
            output_path, pagesize=A4,
            leftMargin=15 * mm, rightMargin=15 * mm,
            topMargin=20 * mm, bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        try:
            korean_font = "Korean"
            title_style = ParagraphStyle("KorTitle", parent=styles["Title"], fontName=korean_font, fontSize=18)
            heading_style = ParagraphStyle("KorHeading", parent=styles["Heading2"], fontName=korean_font, fontSize=13)
            normal_style = ParagraphStyle("KorNormal", parent=styles["Normal"], fontName=korean_font, fontSize=9)
        except Exception:
            title_style = styles["Title"]
            heading_style = styles["Heading2"]
            normal_style = styles["Normal"]

        elements = []

        # 제목
        elements.append(Paragraph("출장 비용 지출 보고서", title_style))
        elements.append(Spacer(1, 10 * mm))

        # 개요 테이블
        summary = calculate_summary(
            receipts,
            trip_info.get("trip_type", "domestic"),
            trip_info.get("exchange_rate", 1.0)
        )

        elements.append(Paragraph("1. 출장 개요", heading_style))
        overview_data = [
            ["출장자", trip_info.get("employee_name", ""), "소속", trip_info.get("department", "")],
            ["출장 기간", f"{trip_info.get('start_date', '')} ~ {trip_info.get('end_date', '')}", "출장지", trip_info.get("destination", "")],
            ["출장 목적", trip_info.get("purpose", ""), "총 비용", _format_amount(summary["total_amount"])],
        ]
        font_for_table = "Korean" if self.font_registered else "Helvetica"
        t = Table(overview_data, colWidths=[30 * mm, 55 * mm, 25 * mm, 55 * mm])
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_for_table),
            ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.83, 0.89, 0.94)),
            ("BACKGROUND", (2, 0), (2, -1), colors.Color(0.83, 0.89, 0.94)),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 8 * mm))

        # 비용 상세
        elements.append(Paragraph("2. 비용 지출 상세", heading_style))

        is_overseas = trip_info.get("trip_type") == "overseas"
        exchange_rate = trip_info.get("exchange_rate", 1.0)

        if is_overseas:
            detail_header = ["No", "날짜", "카테고리", "사용처", "현지금액", "환율", "원화금액"]
        else:
            detail_header = ["No", "날짜", "카테고리", "사용처", "공급가액", "부가세", "합계"]

        detail_data = [detail_header]
        total = 0
        for i, r in enumerate(receipts):
            amount = r.get("total_amount", 0) or 0
            krw = amount if r.get("currency", "KRW") == "KRW" else round(amount * exchange_rate)
            total += krw
            cat_label = CATEGORY_LABELS.get(r.get("category", "O99"), "기타")

            if is_overseas:
                detail_data.append([
                    str(i + 1), r.get("date", ""), cat_label, r.get("vendor_name", ""),
                    f"{amount:,.0f}", f"{exchange_rate:,.2f}", f"{krw:,.0f}"
                ])
            else:
                detail_data.append([
                    str(i + 1), r.get("date", ""), cat_label, r.get("vendor_name", ""),
                    f"{r.get('supply_amount', 0) or 0:,.0f}",
                    f"{r.get('vat_amount', 0) or 0:,.0f}",
                    f"{amount:,.0f}"
                ])

        detail_data.append(["", "", "", "합계", "", "", f"{total:,.0f}"])

        col_widths = [12 * mm, 22 * mm, 28 * mm, 35 * mm, 22 * mm, 22 * mm, 25 * mm]
        t2 = Table(detail_data, colWidths=col_widths)
        t2.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_for_table),
            ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.12, 0.31, 0.47)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (4, 1), (-1, -1), "RIGHT"),
            ("FONTSIZE", (0, -1), (-1, -1), 9),
            ("BACKGROUND", (0, -1), (-1, -1), colors.Color(0.95, 0.95, 0.95)),
        ]))
        elements.append(t2)
        elements.append(Spacer(1, 8 * mm))

        # 카테고리 요약 (Issue #3 수정)
        elements.append(Paragraph("3. 카테고리별 요약", heading_style))
        cat_data = [["카테고리", "금액", "비율"]]
        for cat, amt in sorted(summary["category_totals"].items()):
            pct = round(amt / max(summary["total_amount"], 1) * 100, 1)
            cat_data.append([cat, f"{amt:,.0f}", f"{pct}%"])
        cat_data.append(["합계", f"{summary['total_amount']:,.0f}", "100%"])

        t3 = Table(cat_data, colWidths=[50 * mm, 40 * mm, 30 * mm])
        t3.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_for_table),
            ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.12, 0.31, 0.47)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
            ("BACKGROUND", (0, -1), (-1, -1), colors.Color(0.95, 0.95, 0.95)),
        ]))
        elements.append(t3)
        elements.append(Spacer(1, 8 * mm))

        # 세무 검증 요약
        trip_label = "해외 출장 (적격증빙 면제)" if summary["trip_type"] == "overseas" else "국내 출장"
        tax_data = [
            ["출장 유형", trip_label],
            ["적격증빙 수취율", f"{summary['qualified_count']}/{summary['total_count']}건 ({summary['qualified_rate']}%)"],
            ["부가세 공제 가능액", _format_amount(summary["vat_deductible"])],
            ["부가세 불공제액", _format_amount(summary["vat_non_deductible"])],
            ["가산세 위험 건", f"{summary['risk_items_count']}건"],
        ]
        t4 = Table(tax_data, colWidths=[50 * mm, 80 * mm])
        t4.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_for_table),
            ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.83, 0.89, 0.94)),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
        ]))
        elements.append(t4)
        elements.append(Spacer(1, 10 * mm))

        # 결재란 (Issue #4 수정)
        elements.append(Paragraph("4. 결재", heading_style))
        approval_data = [
            ["구분", "작성자", "팀장", "재무팀", "최종 승인"],
            ["성명", trip_info.get("employee_name", ""), "", "", ""],
            ["서명", "", "", "", ""],
            ["날짜", datetime.now().strftime("%Y-%m-%d"), "", "", ""],
        ]
        t5 = Table(approval_data, colWidths=[25 * mm, 35 * mm, 35 * mm, 35 * mm, 35 * mm])
        t5.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_for_table),
            ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.12, 0.31, 0.47)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BACKGROUND", (0, 1), (0, -1), colors.Color(0.83, 0.89, 0.94)),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        # 서명 행 높이를 위한 spacer row
        elements.append(t5)

        doc.build(elements)
        return output_path
